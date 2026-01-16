from __future__ import annotations

from io import BytesIO
from typing import Iterable

from docx import Document
from fastapi import APIRouter, File, HTTPException, UploadFile

from backend.translation_service import translate_document

router = APIRouter()


def _iter_paragraph_texts(document: Document) -> Iterable[str]:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text


def _iter_table_texts(document: Document) -> Iterable[str]:
    for table in document.tables:
        row_texts = []
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells]
            row_texts.append("\t".join(cell_texts).strip())
        table_text = "\n".join(text for text in row_texts if text)
        if table_text:
            yield table_text


def _extract_blocks(docx_bytes: bytes) -> list[dict]:
    document = Document(BytesIO(docx_bytes))
    blocks: list[dict] = []
    counter = 1
    for text in _iter_paragraph_texts(document):
        blocks.append({"id": f"p{counter}", "source_text": text})
        counter += 1
    for text in _iter_table_texts(document):
        blocks.append({"id": f"t{counter}", "source_text": text})
        counter += 1
    return blocks


@router.post("/translate")
async def translate(file: UploadFile = File(...)) -> dict:
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="只支援 .docx 檔案")

    try:
        docx_bytes = await file.read()
        blocks = _extract_blocks(docx_bytes)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="DOCX 檔案無效") from exc

    return translate_document(blocks)
