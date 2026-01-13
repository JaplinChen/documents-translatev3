import json
import os
import tempfile
from io import BytesIO
from typing import Iterable

from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, Response, UploadFile
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware

from backend.contracts import coerce_blocks
from backend.services.language_detect import detect_document_languages, detect_language
from backend.services.llm_models import (
    list_gemini_models,
    list_ollama_models,
    list_openai_models,
)
from backend.services.translation_memory import (
    delete_glossary,
    delete_tm,
    get_glossary,
    get_tm,
    seed_glossary,
    seed_tm,
    upsert_glossary,
    upsert_tm,
)
from backend.services.pptx_apply import (
    apply_bilingual,
    apply_chinese_corrections,
    apply_translations,
)
from backend.services.pptx_extract import extract_blocks as extract_pptx_blocks
from backend.services.translate_llm import translate_blocks as translate_pptx_blocks
from backend.translation_service import translate_document

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5193", "http://127.0.0.1:5193"],
    allow_methods=["*"],
    allow_headers=["*"],
)


class GlossaryEntry(BaseModel):
    source_lang: str
    target_lang: str
    source_text: str
    target_text: str
    priority: int | None = 0


class GlossaryDelete(BaseModel):
    id: int


class MemoryEntry(BaseModel):
    source_lang: str
    target_lang: str
    source_text: str
    target_text: str


class MemoryDelete(BaseModel):
    id: int


def iter_paragraph_texts(document: Document) -> Iterable[str]:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text


def iter_table_texts(document: Document) -> Iterable[str]:
    for table in document.tables:
        row_texts = []
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells]
            row_texts.append("\t".join(cell_texts).strip())
        table_text = "\n".join(text for text in row_texts if text)
        if table_text:
            yield table_text


def extract_blocks(docx_bytes: bytes) -> list[dict]:
    document = Document(BytesIO(docx_bytes))
    blocks: list[dict] = []
    counter = 1
    for text in iter_paragraph_texts(document):
        blocks.append({"id": f"p{counter}", "source_text": text})
        counter += 1
    for text in iter_table_texts(document):
        blocks.append({"id": f"t{counter}", "source_text": text})
        counter += 1
    return blocks


@app.post("/translate")
async def translate(file: UploadFile = File(...)) -> dict:
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    try:
        docx_bytes = await file.read()
        blocks = extract_blocks(docx_bytes)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Invalid DOCX file") from exc

    return translate_document(blocks)


SUPPORTED_EXTENSIONS = {".pptx", ".docx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}


def _get_file_extension(filename: str | None) -> str:
    """取得檔案副檔名"""
    if not filename or "." not in filename:
        return ""
    return "." + filename.lower().split(".")[-1]


def _validate_file_type(filename: str | None) -> tuple[bool, str]:
    """驗證檔案類型，回傳 (是否有效, 錯誤訊息)"""
    if not filename:
        return False, "請選擇檔案"
    
    ext = _get_file_extension(filename)
    
    if ext in IMAGE_EXTENSIONS:
        return False, f"不支援圖片檔案 ({filename})。此工具僅支援翻譯 PPTX 和 DOCX 文件中的文字內容。如需翻譯圖片中的文字，請使用支援視覺模型的 LLM API（如 GPT-4o）。"
    
    if ext not in SUPPORTED_EXTENSIONS:
        return False, f"不支援的檔案格式 ({ext})。僅支援 .pptx 和 .docx 檔案"
    
    return True, ""


@app.post("/api/pptx/extract")
async def pptx_extract(file: UploadFile = File(...)) -> dict:
    valid, error_msg = _validate_file_type(file.filename)
    if not valid:
        raise HTTPException(status_code=400, detail=error_msg)
    
    try:
        pptx_bytes = await file.read()
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Invalid PPTX file") from exc

    with tempfile.TemporaryDirectory() as temp_dir:
        input_path = os.path.join(temp_dir, "input.pptx")
        with open(input_path, "wb") as handle:
            handle.write(pptx_bytes)
        blocks = extract_pptx_blocks(input_path)

    language_summary = detect_document_languages(blocks)
    return {"blocks": blocks, "language_summary": language_summary}


@app.post("/api/tm/seed")
async def tm_seed() -> dict:
    seed_glossary(
        [
            ("vi", "zh-TW", "báo cáo", "報告", 10),
            ("vi", "zh-TW", "đề xuất", "提案", 9),
            ("vi", "zh-TW", "cải thiện", "改善", 8),
        ]
    )
    seed_tm(
        [
            ("vi", "zh-TW", "lợi ích", "效益"),
            ("vi", "zh-TW", "chi phí", "成本"),
            ("vi", "zh-TW", "đề xuất cải thiện", "改善建議"),
        ]
    )
    return {"status": "seeded"}


@app.get("/api/tm/glossary")
async def tm_glossary(limit: int = 200) -> dict:
    return {"items": get_glossary(limit=limit)}


@app.post("/api/tm/glossary")
async def tm_glossary_upsert(entry: GlossaryEntry) -> dict:
    upsert_glossary(entry.model_dump())
    return {"status": "ok"}


@app.delete("/api/tm/glossary")
async def tm_glossary_delete(entry: GlossaryDelete) -> dict:
    deleted = delete_glossary(entry.id)
    return {"deleted": deleted}


@app.get("/api/tm/memory")
async def tm_memory(limit: int = 200) -> dict:
    return {"items": get_tm(limit=limit)}


@app.post("/api/tm/memory")
async def tm_memory_upsert(entry: MemoryEntry) -> dict:
    upsert_tm(entry.model_dump())
    return {"status": "ok"}


@app.delete("/api/tm/memory")
async def tm_memory_delete(entry: MemoryDelete) -> dict:
    deleted = delete_tm(entry.id)
    return {"deleted": deleted}


@app.post("/api/tm/glossary/import")
async def tm_glossary_import(file: UploadFile = File(...)) -> dict:
    content = (await file.read()).decode("utf-8")
    entries = []
    for line in content.splitlines():
        if not line.strip():
            continue
        parts = [part.strip() for part in line.split(",")]
        if len(parts) < 4:
            continue
        source_lang, target_lang, source_text, target_text = parts[:4]
        priority = int(parts[4]) if len(parts) > 4 and parts[4].isdigit() else 0
        entries.append((source_lang, target_lang, source_text, target_text, priority))
    seed_glossary(entries)
    return {"status": "ok", "count": len(entries)}


@app.post("/api/tm/memory/import")
async def tm_memory_import(file: UploadFile = File(...)) -> dict:
    content = (await file.read()).decode("utf-8")
    entries = []
    for line in content.splitlines():
        if not line.strip():
            continue
        parts = [part.strip() for part in line.split(",")]
        if len(parts) < 4:
            continue
        source_lang, target_lang, source_text, target_text = parts[:4]
        entries.append((source_lang, target_lang, source_text, target_text))
    seed_tm(entries)
    return {"status": "ok", "count": len(entries)}


@app.get("/api/tm/glossary/export")
async def tm_glossary_export() -> Response:
    items = get_glossary(limit=1000)
    lines = ["source_lang,target_lang,source_text,target_text,priority"]
    for item in items:
        lines.append(
            f"{item['source_lang']},{item['target_lang']},"
            f"{item['source_text']},{item['target_text']},{item['priority']}"
        )
    return Response(content="\n".join(lines), media_type="text/csv")


@app.get("/api/tm/memory/export")
async def tm_memory_export() -> Response:
    items = get_tm(limit=1000)
    lines = ["source_lang,target_lang,source_text,target_text"]
    for item in items:
        lines.append(
            f"{item['source_lang']},{item['target_lang']},"
            f"{item['source_text']},{item['target_text']}"
        )
    return Response(content="\n".join(lines), media_type="text/csv")


@app.post("/api/pptx/languages")
async def pptx_languages(file: UploadFile = File(...)) -> dict:
    valid, error_msg = _validate_file_type(file.filename)
    if not valid:
        raise HTTPException(status_code=400, detail=error_msg)
    try:
        pptx_bytes = await file.read()
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Invalid PPTX file") from exc

    with tempfile.TemporaryDirectory() as temp_dir:
        input_path = os.path.join(temp_dir, "input.pptx")
        with open(input_path, "wb") as handle:
            handle.write(pptx_bytes)
        blocks = extract_pptx_blocks(input_path)

    language_summary = detect_document_languages(blocks)
    return {"language_summary": language_summary}


@app.post("/api/pptx/apply")
async def pptx_apply(
    file: UploadFile = File(...),
    blocks: str = Form(...),
    mode: str = Form("bilingual"),
    bilingual_layout: str = Form("inline"),
    fill_color: str | None = Form(None),
    text_color: str | None = Form(None),
    line_color: str | None = Form(None),
    line_dash: str | None = Form(None),
) -> Response:
    valid, error_msg = _validate_file_type(file.filename)
    if not valid:
        raise HTTPException(status_code=400, detail=error_msg)
    try:
        pptx_bytes = await file.read()
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Invalid PPTX file") from exc

    try:
        blocks_data = coerce_blocks(json.loads(blocks))
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail="Invalid blocks JSON") from exc
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Invalid blocks data") from exc

    if mode not in {"bilingual", "correction", "translated"}:
        raise HTTPException(status_code=400, detail="Unsupported mode")
    if mode == "bilingual" and bilingual_layout not in {"inline", "auto", "new_slide"}:
        raise HTTPException(status_code=400, detail="Unsupported bilingual layout")

    with tempfile.TemporaryDirectory() as temp_dir:
        input_path = os.path.join(temp_dir, "input.pptx")
        output_path = os.path.join(temp_dir, "output.pptx")
        with open(input_path, "wb") as handle:
            handle.write(pptx_bytes)

        if mode == "bilingual":
            apply_bilingual(input_path, output_path, blocks_data, layout=bilingual_layout)
        elif mode == "translated":
            apply_translations(input_path, output_path, blocks_data, mode="direct")
        else:
            apply_chinese_corrections(
                input_path,
                output_path,
                blocks_data,
                fill_color=fill_color,
                text_color=text_color,
                line_color=line_color,
                line_dash=line_dash,
            )

        with open(output_path, "rb") as handle:
            output_bytes = handle.read()

    return Response(
        content=output_bytes,
        media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        headers={"Content-Disposition": "attachment; filename=output.pptx"},
    )


@app.post("/api/pptx/translate")
async def pptx_translate(
    blocks: str = Form(...),
    source_language: str | None = Form(None),
    secondary_language: str | None = Form(None),
    target_language: str | None = Form(None),
    mode: str = Form("bilingual"),
    use_tm: bool = Form(False),
    provider: str | None = Form(None),
    model: str | None = Form(None),
    api_key: str | None = Form(None),
    base_url: str | None = Form(None),
) -> dict:
    try:
        blocks_data = coerce_blocks(json.loads(blocks))
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail="Invalid blocks JSON") from exc
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Invalid blocks data") from exc

    if not target_language:
        raise HTTPException(status_code=400, detail="Target language is required")

    def _prepare_blocks_for_correction(items: list[dict]) -> list[dict]:
        if not source_language or source_language == "auto":
            return items
        prepared = []
        for block in items:
            text = block.get("source_text", "")
            if not text:
                prepared.append(block)
                continue
            lines = []
            for line in text.splitlines():
                lang = detect_language(line)
                if lang == source_language:
                    lines.append(line)
            prepared_block = dict(block)
            prepared_block["source_text"] = "\n".join(lines)
            prepared.append(prepared_block)
        return prepared

    if source_language:
        os.environ["SOURCE_LANGUAGE"] = source_language
    try:
        translated = translate_pptx_blocks(
            _prepare_blocks_for_correction(blocks_data)
            if mode == "correction"
            else blocks_data,
            target_language,
            use_tm=use_tm,
            provider=provider,
            model=model,
            api_key=api_key,
            base_url=base_url,
        )
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {
        "mode": mode,
        "source_language": source_language,
        "target_language": target_language,
        "blocks": translated.get("blocks", []),
    }


@app.post("/api/llm/models")
async def llm_models(
    provider: str = Form(...),
    api_key: str | None = Form(None),
    base_url: str | None = Form(None),
) -> dict:
    provider = provider.lower()
    try:
        if provider in {"openai", "chatgpt"}:
            if not api_key:
                raise HTTPException(status_code=400, detail="API key is required")
            models = list_openai_models(api_key, base_url or "https://api.openai.com/v1")
        elif provider == "gemini":
            if not api_key:
                raise HTTPException(status_code=400, detail="API key is required")
            models = list_gemini_models(
                api_key, base_url or "https://generativelanguage.googleapis.com/v1beta"
            )
        elif provider == "ollama":
            models = list_ollama_models(base_url or "http://localhost:11434")
        else:
            raise HTTPException(status_code=400, detail="Unsupported provider")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    return {"models": models}
