"""
Multi-format Export Service

Supports exporting translations to various formats:
- Word (DOCX) - bilingual comparison table
- Excel (XLSX) - translation comparison spreadsheet
- Plain Text (TXT) - simple text export
- PDF (PDF) - bilingual comparison table
"""

from __future__ import annotations

import io
from xml.sax.saxutils import escape

def export_to_docx(blocks: list[dict], filename: str = "translation") -> io.BytesIO:
    """
    Export blocks to Word document with bilingual comparison table.
    Requires python-docx.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX export. Install with: pip install python-docx"
        ) from e

    doc = Document()
    doc.add_heading("翻譯對照表", 0)

    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "#"
    header_cells[1].text = "原文"
    header_cells[2].text = "譯文"

    # Data rows
    for idx, block in enumerate(blocks, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = block.get("original_text", "")
        row_cells[2].text = block.get("translated_text", "")

    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_to_xlsx(blocks: list[dict], filename: str = "translation") -> io.BytesIO:
    """
    Export blocks to Excel spreadsheet with translation comparison.
    Requires openpyxl.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError as e:
        raise ImportError(
            "openpyxl is required for Excel export. Install with: pip install openpyxl"
        ) from e

    wb = Workbook()
    ws = wb.active
    ws.title = "翻譯對照表"

    # Headers with styling
    headers = ["#", "原文", "譯文", "投影片", "類型"]
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # Data rows
    for idx, block in enumerate(blocks, 1):
        ws.cell(row=idx + 1, column=1, value=idx)
        ws.cell(row=idx + 1, column=2, value=block.get("original_text", ""))
        ws.cell(row=idx + 1, column=3, value=block.get("translated_text", ""))
        ws.cell(row=idx + 1, column=4, value=block.get("slide_index", 0))
        ws.cell(row=idx + 1, column=5, value=block.get("block_type", ""))

    # Adjust column widths
    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 50
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 12

    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def export_to_txt(blocks: list[dict], include_original: bool = True) -> io.BytesIO:
    """
    Export blocks to plain text file.
    """
    lines = []

    for idx, block in enumerate(blocks, 1):
        original = block.get("original_text", "")
        translated = block.get("translated_text", "")

        if include_original:
            lines.append(f"[{idx}] 原文: {original}")
            lines.append(f"    譯文: {translated}")
            lines.append("")
        else:
            lines.append(f"[{idx}] {translated}")

    content = "\n".join(lines)
    output = io.BytesIO(content.encode("utf-8"))
    output.seek(0)
    return output


def export_to_pdf(blocks: list[dict], filename: str = "translation") -> io.BytesIO:
    """
    匯出翻譯區塊為 PDF 對照表。
    需要 reportlab。
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as e:
        raise ImportError(
            "PDF 匯出需要 reportlab。請安裝: pip install reportlab"
        ) from e

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
    )

    pdfmetrics.registerFont(UnicodeCIDFont("MSung-Light"))

    styles = getSampleStyleSheet()
    base_style = ParagraphStyle(
        "CJKBody",
        parent=styles["Normal"],
        fontName="MSung-Light",
        fontSize=10,
        leading=14,
    )
    title_style = ParagraphStyle(
        "CJKTitle",
        parent=styles["Title"],
        fontName="MSung-Light",
        fontSize=16,
        leading=20,
    )

    elements = [Paragraph("翻譯對照表", title_style), Spacer(1, 12)]

    data = [["#", "原文", "譯文"]]
    for idx, block in enumerate(blocks, 1):
        original = escape(block.get("original_text", "")).replace("\n", "<br/>")
        translated = escape(block.get("translated_text", "")).replace("\n", "<br/>")
        data.append(
            [
                str(idx),
                Paragraph(original, base_style),
                Paragraph(translated, base_style),
            ]
        )

    index_width = 32
    text_width = (doc.width - index_width) / 2
    table = Table(data, colWidths=[index_width, text_width, text_width])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, -1), "MSung-Light"),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("ALIGN", (0, 0), (0, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    elements.append(table)
    doc.build(elements)
    output.seek(0)
    return output


def get_export_formats() -> list[dict]:
    """Get list of available export formats."""
    return [
        {"id": "pptx", "label": "PowerPoint (.pptx)", "icon": "📊", "available": True},
        {"id": "docx", "label": "Word 對照表 (.docx)", "icon": "📝", "available": True},
        {"id": "xlsx", "label": "Excel 對照表 (.xlsx)", "icon": "📈", "available": True},
        {"id": "txt", "label": "純文字 (.txt)", "icon": "📄", "available": True},
        {
            "id": "pdf",
            "label": "PDF 對照表 (.pdf)",
            "icon": "🖨️",
            "available": True,
        },
    ]
