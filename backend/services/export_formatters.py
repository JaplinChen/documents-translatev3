from __future__ import annotations

import io


def _get_original_text(block: dict) -> str:
    return (
        block.get("original_text")
        or block.get("source_text")
        or block.get("text")
        or ""
    )


def _get_translated_text(block: dict) -> str:
    return (
        block.get("translated_text")
        or block.get("target_text")
        or ""
    )


def export_to_docx(
    blocks: list[dict],
    filename: str = "translation",
) -> io.BytesIO:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install with: pip install python-docx"
        ) from e

    doc = Document()
    doc.add_heading("翻譯對照表", 0)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "#"
    header_cells[1].text = "原文"
    header_cells[2].text = "譯文"

    for idx, block in enumerate(blocks, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = _get_original_text(block)
        row_cells[2].text = _get_translated_text(block)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_to_xlsx(
    blocks: list[dict],
    filename: str = "translation",
) -> io.BytesIO:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError as e:
        raise ImportError(
            "openpyxl is required for Excel export. "
            "Install with: pip install openpyxl"
        ) from e

    wb = Workbook()
    ws = wb.active
    ws.title = "翻譯對照表"

    headers = ["#", "原文", "譯文", "投影片", "類型"]
    header_fill = PatternFill(
        start_color="4472C4",
        end_color="4472C4",
        fill_type="solid",
    )
    header_font = Font(color="FFFFFF", bold=True)

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    for idx, block in enumerate(blocks, 1):
        ws.cell(row=idx + 1, column=1, value=idx)
        ws.cell(row=idx + 1, column=2, value=_get_original_text(block))
        ws.cell(row=idx + 1, column=3, value=_get_translated_text(block))
        ws.cell(row=idx + 1, column=4, value=block.get("slide_index", 0))
        ws.cell(row=idx + 1, column=5, value=block.get("block_type", ""))

    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 50
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 12

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def export_to_txt(
    blocks: list[dict],
    include_original: bool = True,
) -> io.BytesIO:
    lines = []

    for idx, block in enumerate(blocks, 1):
        original = _get_original_text(block)
        translated = _get_translated_text(block)

        if include_original:
            lines.append(f"[{idx}] 原文: {original}")
            lines.append(f"    譯文: {translated}")
            lines.append("")
        else:
            lines.append(f"[{idx}] {translated}")

    content = "\n".join(lines)
    output = io.BytesIO(content.encode("utf-8"))
    output.seek(0)
    return output


def export_to_md(
    blocks: list[dict],
) -> io.BytesIO:
    def _md_escape(text: str) -> str:
        return text.replace("|", "\\|").replace("\n", "<br>")

    lines = []
    lines.append("# 翻譯對照表")
    lines.append("")
    lines.append("| # | 原文 | 譯文 |")
    lines.append("| --- | --- | --- |")
    for idx, block in enumerate(blocks, 1):
        original = _md_escape(_get_original_text(block))
        translated = _md_escape(_get_translated_text(block))
        lines.append(f"| {idx} | {original} | {translated} |")

    content = "\n".join(lines)
    output = io.BytesIO(content.encode("utf-8"))
    output.seek(0)
    return output
